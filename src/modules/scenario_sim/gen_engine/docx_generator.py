"""
Module for generating DOCX documents from Scenario instances.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Import the generated Pydantic models
from modules.scenario_sim.gen_engine.gen_parts.scenario import Scenario


class DocxGenerator:
    """
    Generates a DOCX document based on a Scenario model and a template.
    The generated document preserves the headers and footers of the template
    and creates tables with black borders and light-blue header shading where
    applicable.
    """

    def __init__(self, template_path: str = ""):
        """
        Initialise the generator with a template DOCX file.

        Args:
            template_path: Path to the template DOCX file.
        """
        path = self._set_template_path(template_path)
        print(path)
        self.doc = Document(path)

    @staticmethod
    def _set_template_path(template_path: str = "") -> Path:
        # Default template path
        if template_path == "":
            template_path = os.path.join(
                os.path.dirname(__file__),
                "..",
                "..",
                "..",
                "..",
                "assets",
                "Template Cenário_Facens.docx",
            )
            template_path = os.path.abspath(template_path)

        # Check if template exists
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        return Path(template_path)

    def _add_section(self, title: str, rows: List[List[str | None]]) -> None:
        """
        Add a titled section to the document.

        The section is rendered as a table where each row contains elements from the tuples in `rows`.
        All cells will be merged for the title which will be placed at the top of the section.
        Table style is set to "Unifacens".

        Args:
            title: Title to display at the top of the section.
            rows: List of lists, where each inner list represents a row in the table and
                  each element in the inner list represents a cell value for that row.
        """
        # Determine the number of columns based on the length of the first row (assuming all rows have the same length)
        cols_count = len(rows[0]) if rows else 0

        # Create a table with the calculated dimensions
        total_rows = 1 + len(rows)  # one for title and rest for data
        table = self.doc.add_table(rows=total_rows, cols=cols_count)
        table.style = "Unifacens"

        # Populate the title cell (spanning all columns)
        title_cell = table.cell(0, 0)
        title_cells = [table.cell(0, col) for col in range(cols_count)]
        for cell in title_cells:
            if cell != title_cell:
                title_cell.merge(cell)

        title_cell.text = title
        title_cell.paragraphs[0].runs[0].font.size = Pt(14)
        title_cell.paragraphs[0].runs[0].font.bold = True
        title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Populate the data rows
        for i, row in enumerate(rows):
            start_row = 1 + i
            for col_index, value in enumerate(row):
                cell = table.cell(start_row, col_index)
                cell.text = str(value)

    def generate(self, scenario: Scenario, output_path: str) -> str:
        """
        Populate the document with data from a Scenario instance and save it.

        Args:
            scenario: The Scenario model instance containing all data.
            output_path: Path where the generated DOCX should be written.
        """
        # Clear all paragraphs and tables from self.doc
        for para in list(self.doc.paragraphs):
            p = para._element.getparent()
            p.remove(para._element)

        for table in list(self.doc.tables):
            t = table._tbl
            tbl_elm = t.getparent()
            tbl_elm.remove(t)

        # Section: Course Information
        self._add_section(
            "ROTEIRO DE CENÁRIO DE SIMULAÇÃO",
            [
                ["Dia e horário da aula:", None],  # Placeholder: to be filled later
                ["Local de realização da aula:", None],
                ["Nome do cenário:", None],  # Placeholder: to be filled later
                ["Tempo de duração:", None],  # Placeholder: to be filled later
                ["Curso(s):", None],  # Placeholder: to be filled later
                ["Unidade curricular:", None],  # Placeholder: to be filled later
                ["Turma:", None],  # Placeholder: to be filled later
                ["Quantidade de estudantes:", None],  # Placeholder: to be filled later
                ["Professor:", None],  # Placeholder: to be filled later
            ],
        )
        self.doc.add_paragraph()

        # Section: Learning Objectives
        self._add_section("Objetivos de Aprendizagem", [[scenario.learning_objectives]])
        self.doc.add_paragraph()

        # Section: Resources
        self._add_section(
            "Recursos materiais necessários:",
            [[r.name, str(r.quantity)] for r in scenario.necessary_resources],
        )
        self.doc.add_paragraph()

        # Section: Scene Organization
        self._add_section(
            "Organização do ambiente:",
            [[scenario.scene_organization]],
        )
        self.doc.add_paragraph()

        # Section: Participants
        participants = scenario.scene_participants
        uses_sim = "x" if participants.uses_simulator else " "
        self._add_section(
            "Participantes do cenário",
            [
                [
                    "Participante:",
                    f"Estudante ({str(participants.students_quantity)})",
                    f"Ator ({str(participants.actors_quantity)})",
                    f"Simulador ({uses_sim})",
                ],
                [
                    "Função no cenário:",
                    participants.students_role,
                    participants.actors_role,
                    participants.simulator_role,
                ],
            ],
        )
        self.doc.add_paragraph()

        # Section: Case Presentation
        self._add_section("Apresentação do caso", [[scenario.case_presentation]])
        self.doc.add_paragraph()

        # Section: Actor Briefings
        for idx, actor in enumerate(scenario.actor_briefing, start=1):
            self._add_section(
                f"Briefing do ator {idx}",
                [
                    ["Dados pessoais:", actor.personal_data],
                    ["História atual:", actor.current_story],
                    ["História prévia:", actor.previous_story],
                    ["Vestimentas:", actor.clothing],
                    ["Perfil de comportamento:", actor.behavior_profile],
                ],
            )
            self.doc.add_paragraph()

        # Section: Simulator Parameters
        self._add_section(
            "Programação do Simulador:", [[scenario.simulator_parameters]]
        )
        self.doc.add_paragraph()

        # Section: Evolution Parameters
        self._add_section(
            "Programação evolutivos do simulador:",
            [[scenario.simulator_evolution_parameters]],
        )
        self.doc.add_paragraph()

        # Section: Students Briefing
        self._add_section(
            "Briefing para os estudantes:", [[scenario.students_briefing]]
        )
        self.doc.add_paragraph()

        # Section: Scene Flow
        for idx, scene in enumerate(scenario.scene_flow, start=1):
            self._add_section(
                f"Cena {idx}",
                [
                    ["Plano A:", scene.student_plan_a],
                    ["Ator:", scene.actor_sim_directions],
                    ["Plano B:", scene.actor_plan_b],
                ],
            )
            self.doc.add_paragraph()

        # Section: Debriefing
        self._add_section("Debriefing:", [[scenario.debriefing]])
        self.doc.add_paragraph()

        # Section: Appendix
        self._add_section("Anexos:", [[scenario.appendix]])

        # Save the final document
        self.doc.save(output_path)
        return output_path
