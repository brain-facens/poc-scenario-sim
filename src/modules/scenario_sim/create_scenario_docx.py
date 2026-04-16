#!/usr/bin/env python3
"""
Script to generate DOCX files from Scenario objects using a template.
This script creates a document based on the Scenario model structure while
preserving the template's header and footer and following ABNT formatting
standards. No tables are used in the final document; all data is presented
using paragraphs and bullet points for readability.
"""

import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from modules.scenario_sim.gen_engine.gen_parts.actor_briefing import ActorBriefing
from modules.scenario_sim.gen_engine.gen_parts.participants import Participants
from modules.scenario_sim.gen_engine.gen_parts.resource import Resource
from modules.scenario_sim.gen_engine.gen_parts.scenario import Scenario
from modules.scenario_sim.gen_engine.gen_parts.scene import Scene


class DocxExporter:
    def __init__(self, template_path: str | None = None):
        """
        Initialise the exporter with a template DOCX file.
        If template_path is None, a default path relative to this file is used.
        """
        self._set_template_path(template_path)
        self.doc = Document(self.template_path)

    def _set_template_path(self, template_path: str | None):
        """
        Determine the path to the template DOCX file.
        """
        if template_path is None:
            # Default template path located in the assets folder of the project root.
            base_dir = os.path.abspath(
                os.path.join(os.path.dirname(__file__), "..", "..", "..")
            )
            template_path = os.path.join(
                base_dir, "assets", "Template Cenário_Facens.docx"
            )
            template_path = os.path.abspath(template_path)

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template file not found: {template_path}")

        self.template_path = template_path

    def create_scenario_docx(self, scenario: Scenario, output_path: str) -> str:
        """
        Generate a DOCX document from the given scenario instance.
        The generated document preserves the template's headers and footers
        and uses paragraphs (no tables) to display all information.
        """
        doc = self.doc

        # Set document properties (metadata)
        doc.core_properties.title = "Cenário de Simulação"
        doc.core_properties.author = "UniFacens"
        doc.core_properties.created = datetime.now()

        # Preserve header and footer: remove all intermediate paragraphs
        # but keep the first (header) and the last (footer) paragraph.
        num_paragraphs = len(doc.paragraphs)
        for i in range(num_paragraphs - 2, 0, -1):
            p = doc.paragraphs[i]
            p._p.getparent().remove(p._p)

        # Remove any tables (tables in headers/footers may be desired,
        # but removing them here simplifies the final document).
        for i in range(len(doc.tables) - 1, -1, -1):
            tbl = doc.tables[i]
            tbl._tbl.getparent().remove(tbl._tbl)

        # Add title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("UniFacens - Centro Universitário Facens")
        title_run.bold = True
        title_run.font.size = Pt(24)

        title2 = doc.add_paragraph()
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title2_run = title2.add_run("Cenário de Simulação")
        title2_run.bold = True
        title2_run.font.size = Pt(24)

        doc.add_paragraph()  # Blank line

        # Main information (previously a table)
        main_fields = [
            ("Dia e horário da aula", ""),
            ("Local de realização da aula", ""),
            ("Nome do cenário", ""),
            ("Tempo de duração", ""),
            ("Curso(s)", ""),
            ("Unidade curricular", ""),
            ("Turma", ""),
            ("Quantidade de estudantes", ""),
            ("Professor", ""),
        ]
        for label, value in main_fields:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"{label}: {value}")
            run.bold = True
            run.font.size = Pt(12)

        doc.add_paragraph()  # Blank line

        # Learning objectives
        obj_title = doc.add_paragraph()
        obj_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        obj_title_run = obj_title.add_run("Objetivos de Aprendizagem:")
        obj_title_run.bold = True
        obj_title_run.font.size = Pt(14)

        obj_run = doc.add_paragraph()
        obj_run.add_run(scenario.learning_objectives).font.size = Pt(12)

        doc.add_paragraph()  # Blank line

        # Necessary resources
        res_title = doc.add_paragraph()
        res_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        res_title_run = res_title.add_run("Recursos Necessários")
        res_title_run.bold = True
        res_title_run.font.size = Pt(14)

        if scenario.necessary_resources:
            for res in scenario.necessary_resources:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.add_run(f"{res.name}: {res.quantity}x")
        else:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("Nenhum recurso necessário").italic = True

        doc.add_paragraph()  # Blank line

        # Scene organization
        org_title = doc.add_paragraph()
        org_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        org_title_run = org_title.add_run("Organização da Cena")
        org_title_run.bold = True
        org_title_run.font.size = Pt(14)

        org_run = doc.add_paragraph()
        org_run.add_run(scenario.scene_organization).font.size = Pt(12)

        doc.add_paragraph()  # Blank line

        # Participants
        parts_title = doc.add_paragraph()
        parts_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        parts_title_run = parts_title.add_run("Participantes do cenário")
        parts_title_run.bold = True
        parts_title_run.font.size = Pt(14)

        participants_info = [
            f"Estudante ({scenario.scene_participants.students_quantity}) - Função: {scenario.scene_participants.students_role}",
            f"Ator ({scenario.scene_participants.actors_quantity}) - Função: {scenario.scene_participants.actors_role}",
            f"Simulador ({int(scenario.scene_participants.uses_simulator)}) - Função: {scenario.scene_participants.simulator_role}",
        ]
        for info in participants_info:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.add_run(info)
        doc.add_paragraph()  # Blank line

        # Case presentation
        case_title = doc.add_paragraph()
        case_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        case_title_run = case_title.add_run("Apresentação do Caso")
        case_title_run.bold = True
        case_title_run.font.size = Pt(14)

        case_run = doc.add_paragraph()
        case_run.add_run(scenario.case_presentation).font.size = Pt(12)
        doc.add_paragraph()  # Blank line

        # Actor briefings
        for idx, actor in enumerate(scenario.actor_briefing, 1):
            brief_title = doc.add_paragraph()
            brief_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            brief_title_run = brief_title.add_run(f"Briefing do ator {idx}")
            brief_title_run.bold = True
            brief_title_run.font.size = Pt(14)

            # Personal data
            p = doc.add_paragraph()
            p.add_run("Dados Pessoais:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(actor.personal_data)
            p = doc.add_paragraph()
            # Current story
            p = doc.add_paragraph()
            p.add_run("História atual:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(actor.current_story)
            p = doc.add_paragraph()
            # Previous story
            p = doc.add_paragraph()
            p.add_run("História anterior:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(actor.previous_story)
            p = doc.add_paragraph()
            # Clothing
            p = doc.add_paragraph()
            p.add_run("Vestimenta:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(actor.clothing)
            p = doc.add_paragraph()
            # Behavior profile
            p = doc.add_paragraph()
            p.add_run("Perfil de comportamento:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(actor.behavior_profile)
            doc.add_paragraph()  # Blank line after each actor

        doc.add_paragraph()  # Blank line

        # Simulator parameters
        sim_title = doc.add_paragraph()
        sim_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sim_title_run = sim_title.add_run("Parâmetros do Simulador")
        sim_title_run.bold = True
        sim_title_run.font.size = Pt(14)

        sim_run = doc.add_paragraph()
        sim_run.add_run(scenario.simulator_parameters).font.size = Pt(12)

        doc.add_paragraph()  # Blank line

        # Evolution parameters
        ev_title = doc.add_paragraph()
        ev_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        ev_title_run = ev_title.add_run("Parâmetros de Evolução do Simulador")
        ev_title_run.bold = True
        ev_title_run.font.size = Pt(14)

        ev_run = doc.add_paragraph()
        ev_run.add_run(scenario.simulator_evolution_parameters).font.size = Pt(12)

        doc.add_paragraph()  # Blank line

        # Students briefing
        stud_title = doc.add_paragraph()
        stud_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        stud_title_run = stud_title.add_run("Briefing dos Estudantes")
        stud_title_run.bold = True
        stud_title_run.font.size = Pt(14)

        stud_run = doc.add_paragraph()
        stud_run.add_run(scenario.students_briefing).font.size = Pt(12)

        doc.add_paragraph()  # Blank line

        # Scene flow
        for idx, scene in enumerate(scenario.scene_flow, 1):
            flow_title = doc.add_paragraph()
            flow_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            flow_title_run = flow_title.add_run(f"Cena {idx}")
            flow_title_run.bold = True
            flow_title_run.font.size = Pt(14)

            # Student plan (ESTUDANTE)
            p = doc.add_paragraph()
            p.add_run("ESTUDANTE:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(scene.student_plan_a)
            p = doc.add_paragraph()

            # Actor/Simulator directions (ATOR/SIMULADOR)
            p = doc.add_paragraph()
            p.add_run("ATOR/SIMULADOR:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(scene.actor_sim_directions)
            p = doc.add_paragraph()

            # Plan B (PLANO B)
            p = doc.add_paragraph()
            p.add_run("PLANO B:")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.add_run(scene.actor_plan_b)

            doc.add_paragraph()  # Blank line after each scene

        doc.add_paragraph()  # Blank line

        # Debriefing
        debrief_title = doc.add_paragraph()
        debrief_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        debrief_title_run = debrief_title.add_run("Debriefing")
        debrief_title_run.bold = True
        debrief_title_run.font.size = Pt(14)

        debrief_run = doc.add_paragraph()
        debrief_run.add_run(scenario.debriefing).font.size = Pt(12)

        doc.add_paragraph()  # Blank line

        # Appendix
        appendix_title = doc.add_paragraph()
        appendix_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        appendix_title_run = appendix_title.add_run("Anexo")
        appendix_title_run.bold = True
        appendix_title_run.font.size = Pt(14)

        appendix_run = doc.add_paragraph()
        appendix_run.add_run(scenario.appendix).font.size = Pt(12)

        # Save the document
        doc.save(output_path)
        print(f"Document saved to {output_path}")
        return output_path


if __name__ == "__main__":
    # Example usage (for debugging purposes only)
    pass
